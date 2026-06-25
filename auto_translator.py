# -*- coding: utf-8 -*-
"""
자동 언어감지 -> 한국어 번역기 (로컬 웹 UI)
- 실행하면 브라우저가 열립니다.
- 파일 경로를 입력하면 언어를 자동 감지해 한국어로 번역하고
  같은 폴더에 "<원본이름>_KO.<확장자>" 로 저장합니다.
- 지원 형식: .docx  .txt  (서식/이미지/표 그대로 유지)
필요 패키지: lxml (docx 처리용)
"""
import os, re, sys, json, time, html, threading, uuid, webbrowser
import urllib.request, urllib.parse
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from concurrent.futures import ThreadPoolExecutor

PORT = 8765
W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
JOBS = {}                 # job_id -> 진행상태 dict
JOBS_LOCK = threading.Lock()

# 장식용 아이콘(번역 대상 아님)으로 쓰이는 기호 글꼴들
SYMBOL_FONT_RE = re.compile(r'wingding|webding|symbol|marlett|dingbat', re.I)
# 번역할 가치가 있는 '글자'(라틴/그리스/키릴/CJK/가나 등)
LETTER_RE = re.compile(
    r'[A-Za-zÀ-ɏͰ-ϿЀ-ӿ'
    r'一-鿿぀-ヿ]')


def is_symbol_run(r):
    """Wingdings 등 기호 글꼴이거나 <w:sym> 아이콘을 쓰는 런이면 True (번역 제외)."""
    rpr = r.find(W + 'rPr')
    if rpr is None:
        return False
    if rpr.find(W + 'sym') is not None:
        return True
    rf = rpr.find(W + 'rFonts')
    if rf is not None:
        fonts = ' '.join(filter(None, [
            rf.get(W + 'ascii'), rf.get(W + 'hAnsi'),
            rf.get(W + 'cs'), rf.get(W + 'eastAsia')]))
        if SYMBOL_FONT_RE.search(fonts):
            return True
    return False


# ----------------------------- 번역 엔진 -----------------------------
def google_translate(text, sl='auto', tl='ko'):
    """구글 무료 엔드포인트로 한 덩어리 번역. (번역문, 감지언어) 반환."""
    url = ('https://translate.googleapis.com/translate_a/single'
           '?client=gtx&sl=%s&tl=%s&dt=t&q=%s'
           % (sl, tl, urllib.parse.quote(text)))
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    last = None
    for attempt in range(6):
        try:
            r = urllib.request.urlopen(req, timeout=20)
            data = json.loads(r.read().decode('utf-8'))
            out = ''.join(seg[0] for seg in data[0] if seg and seg[0])
            detected = data[2] if len(data) > 2 else sl
            return out, detected
        except Exception as e:
            last = e
            time.sleep(1.5 * (attempt + 1))
    raise last


def translate_unique(strings, job, max_workers=8):
    """중복 제거된 문자열들을 병렬 번역. cache(dict)와 감지언어를 반환."""
    uniq = sorted(set(strings))
    cache, lock, done, detected = {}, threading.Lock(), [0], {'lang': None}

    def work(s):
        try:
            out, lang = google_translate(s)
        except Exception:
            out, lang = s, None     # 실패 시 원문 유지
        with lock:
            cache[s] = out
            if lang and not detected['lang']:
                detected['lang'] = lang
            done[0] += 1
            job['done'] = done[0]

    job['total'] = len(uniq)
    with ThreadPoolExecutor(max_workers=max_workers) as ex:
        list(ex.map(work, uniq))
    return cache, detected['lang']


# ----------------------------- 형식별 처리 -----------------------------
def translate_docx(src_path, out_path, job):
    from lxml import etree
    import zipfile, shutil, tempfile

    tmp = tempfile.mkdtemp(prefix='trdocx_')
    try:
        with zipfile.ZipFile(src_path) as z:
            z.extractall(tmp)
        doc_xml = os.path.join(tmp, 'word', 'document.xml')
        tree = etree.parse(doc_xml)
        root = tree.getroot()

        # 문단 단위로 묶어 번역 (단어가 run 단위로 쪼개지는 문제 방지)
        # 단, Wingdings 등 '기호 글꼴' 런은 번역 대상에서 제외하고 그대로 둔다.
        # (그렇지 않으면 번역문이 아이콘 글꼴 런에 들어가 깨져 보임)
        para_runs, sources = [], []
        for p in root.iter(W + 'p'):
            ts = []
            for r in p.iter(W + 'r'):
                anc = r.getparent()
                while anc is not None and anc.tag != W + 'p':
                    anc = anc.getparent()
                if anc is not p:
                    continue
                if is_symbol_run(r):       # 아이콘/기호 런은 손대지 않음
                    continue
                t = r.find(W + 't')
                if t is not None:
                    ts.append(t)
            s = ''.join((t.text or '') for t in ts)
            if LETTER_RE.search(s):
                para_runs.append(ts)
                sources.append(s)

        job['stage'] = '번역 중'
        cache, lang = translate_unique(sources, job)

        job['stage'] = '문서에 반영 중'
        space = '{http://www.w3.org/XML/1998/namespace}space'
        for ts, s in zip(para_runs, sources):
            ts[0].text = cache.get(s, s)
            ts[0].set(space, 'preserve')
            for t in ts[1:]:
                t.text = ''
        new_data = etree.tostring(tree, xml_declaration=True,
                                  encoding='UTF-8', standalone=True)

        # 원본 zip의 모든 항목 복사 + document.xml만 교체
        with zipfile.ZipFile(src_path) as zin, \
             zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/document.xml':
                    zout.writestr(item, new_data)
                else:
                    zout.writestr(item, zin.read(item.filename))
        return lang
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def translate_txt(src_path, out_path, job):
    raw = open(src_path, 'r', encoding='utf-8', errors='replace').read()
    # 빈 줄 기준으로 문단 분리, 빈 문단/공백은 그대로 둠
    parts = re.split(r'(\n\s*\n)', raw)
    targets = [p for p in parts if p.strip() and not re.fullmatch(r'\n\s*\n', p)]
    job['stage'] = '번역 중'
    cache, lang = translate_unique(targets, job)
    out = ''.join(cache.get(p, p) if (p.strip() and not re.fullmatch(r'\n\s*\n', p)) else p
                  for p in parts)
    open(out_path, 'w', encoding='utf-8').write(out)
    return lang


def run_job(job_id, path, formats=None):
    job = JOBS[job_id]
    try:
        path = path.strip().strip('"').strip("'")
        formats = [f.lower().strip() for f in (formats or ['docx']) if f and f.strip()]
        if not formats:
            formats = ['docx']
        want = lambda f: f in formats

        # ---- 1) 번역된 '기준 DOCX' 생성 ----
        if re.match(r'(?i)^https?://', path):
            job['stage'] = '웹페이지 가져오는 중'
            base_docx, lang, out_dir, stem = translate_url(path, job)
        else:
            if not os.path.isfile(path):
                raise FileNotFoundError('파일 또는 주소를 찾을 수 없습니다: %s' % path)
            out_dir = os.path.dirname(path)
            stem, ext = os.path.splitext(os.path.basename(path))
            ext_low = ext.lower()
            job['stage'] = '파일 분석 중'

            # 스프레드시트 입력 -> 번역 엑셀(서식·수식 보존)로 바로 저장하고 종료
            if ext_low in ('.xlsx', '.xlsm', '.xls'):
                out_xlsx = os.path.join(out_dir, stem + '_KO.xlsx')
                if ext_low == '.xls':
                    lang = translate_xls_to_xlsx(path, out_xlsx, job)
                else:
                    lang = translate_xlsx(path, out_xlsx, job)
                job['status'] = 'done'
                job['stage'] = '완료'
                job['output'] = out_xlsx
                job['lang'] = lang or '자동감지'
                return

            base_docx = os.path.join(out_dir, stem + '_KO.docx')
            if ext_low == '.docx':
                lang = translate_docx(path, base_docx, job)
            elif ext_low == '.doc':
                lang = translate_doc_to_docx(path, base_docx, job)
            elif ext_low == '.txt':
                lang = txt_to_docx(path, base_docx, job)
            elif ext_low == '.pdf':
                lang = translate_pdf_to_docx(path, base_docx, job)
            else:
                raise ValueError('지원하지 않는 형식입니다: %s\n'
                                 '(지원: docx, doc, pdf, txt, xlsx, xls, 웹주소)' % ext)

        # ---- 2) 선택한 형식들로 변환/저장 ----
        base_abs = os.path.abspath(base_docx)
        outputs, warns = [], []

        if want('docx'):
            outputs.append(base_docx)

        need_pdf = want('pdf')
        need_jpg = want('jpeg') or want('jpg')
        if need_pdf or need_jpg:
            pdf_path = os.path.join(out_dir, stem + '_KO.pdf')
            pdf_ok = False
            job['stage'] = 'PDF 변환 중'
            try:
                docx_to_pdf(base_abs, os.path.abspath(pdf_path))
                pdf_ok = os.path.exists(pdf_path)
            except Exception as e:
                warns.append('PDF 생성 실패(MS Word 필요): %s' % e)

            if need_pdf and pdf_ok:
                outputs.append(pdf_path)
                if not _is_real_pdf(pdf_path):
                    warns.append('PDF는 회사 DRM(Fasoo)으로 보호되어 인가된 앱에서만 열립니다.')

            if need_jpg:
                job['stage'] = 'JPEG 이미지 생성 중'
                jdir = os.path.join(out_dir, stem + '_KO_jpg')
                imgs = []
                if pdf_ok and _is_real_pdf(pdf_path):
                    try:
                        imgs = pdf_to_jpegs(pdf_path, jdir, stem)
                    except Exception:
                        imgs = []
                if not imgs:   # DRM 등으로 페이지 캡처 불가 시 텍스트 렌더 폴백
                    imgs = render_text_to_jpegs(docx_text_lines(base_docx), jdir, stem)
                    warns.append('JPEG는 DRM으로 페이지 캡처가 불가해 번역 텍스트 이미지로 생성했습니다.')
                if imgs:
                    outputs.append('%s  (이미지 %d장)' % (jdir, len(imgs)))

            if not need_pdf and os.path.exists(pdf_path):
                os.remove(pdf_path)

        if want('excel') or want('xlsx'):
            job['stage'] = 'Excel 작성 중'
            xlsx_path = os.path.join(out_dir, stem + '_KO.xlsx')
            try:
                docx_tables_to_excel(base_abs, xlsx_path)
                outputs.append(xlsx_path)
            except Exception as e:
                warns.append('Excel 변환 실패: %s' % e)

        # docx를 원치 않으면 기준 파일 삭제 (변환에만 사용)
        if not want('docx') and os.path.exists(base_docx):
            os.remove(base_docx)

        job['status'] = 'done'
        job['stage'] = '완료'
        msg = '\n'.join(outputs) if outputs else '(생성된 파일 없음)'
        if warns:
            msg += '\n⚠ ' + '\n⚠ '.join(warns)
        job['output'] = msg
        job['lang'] = lang or '자동감지'
    except Exception as e:
        job['status'] = 'error'
        job['stage'] = '오류'
        job['error'] = str(e)


def translate_url(url, job):
    """웹페이지를 가져와 본문을 한국어로 번역하고 docx로 저장. (out_path, lang) 반환."""
    import docx
    from bs4 import BeautifulSoup

    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    raw = urllib.request.urlopen(req, timeout=30).read()
    soup = BeautifulSoup(raw, 'lxml')

    # 제목
    title = (soup.title.string.strip() if soup.title and soup.title.string
             else urllib.parse.urlparse(url).netloc)

    # 본문에서 필요 없는 요소 제거
    for tag in soup(['script', 'style', 'noscript', 'head', 'svg',
                     'nav', 'footer', 'header', 'form', 'aside', 'iframe']):
        tag.decompose()

    # 블록 단위로 (스타일, 텍스트) 수집
    blocks = []
    for el in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                             'p', 'li', 'blockquote', 'td', 'pre']):
        txt = ' '.join(el.get_text(' ', strip=True).split())
        if not txt or not LETTER_RE.search(txt):
            continue
        if el.name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            style = ('H', int(el.name[1]))
        else:
            style = ('P', 0)
        blocks.append((style, txt))

    if not blocks:
        raise ValueError('웹페이지에서 번역할 본문 텍스트를 찾지 못했습니다.')

    job['stage'] = '번역 중'
    cache, lang = translate_unique([t for _, t in blocks], job)

    job['stage'] = '문서 작성 중'
    doc = docx.Document()
    doc.add_heading(cache.get(title, title) if LETTER_RE.search(title) else title, 0)
    p = doc.add_paragraph()
    run = p.add_run('원본: ' + url)
    run.italic = True
    for (kind, level), txt in blocks:
        tr = cache.get(txt, txt)
        if kind == 'H':
            doc.add_heading(tr, min(level, 4))
        else:
            doc.add_paragraph(tr)

    # 저장 위치: 프로그램 폴더 (URL은 원본 폴더가 없으므로)
    out_dir = os.path.dirname(os.path.abspath(__file__))
    base = re.sub(r'[\\/:*?"<>|]+', '_', title)[:60].strip() or 'webpage'
    stem, n = base, 1
    while os.path.exists(os.path.join(out_dir, stem + '_KO.docx')):
        n += 1
        stem = '%s(%d)' % (base, n)
    out_path = os.path.join(out_dir, stem + '_KO.docx')
    doc.save(out_path)
    return out_path, lang, out_dir, stem


def txt_to_docx(src_path, out_docx, job):
    """텍스트 파일을 번역해 docx로 만든다. 감지언어 반환."""
    import docx as _docx
    raw = open(src_path, 'r', encoding='utf-8', errors='replace').read()
    parts = re.split(r'\n\s*\n', raw)
    targets = [p for p in parts if p.strip()]
    job['stage'] = '번역 중'
    cache, lang = translate_unique(targets, job)
    job['stage'] = '문서 작성 중'
    doc = _docx.Document()
    for p in parts:
        if p.strip():
            doc.add_paragraph(cache.get(p, p))
    doc.save(out_docx)
    return lang


def translate_doc_to_docx(path, out_docx, job):
    """구형 .doc 를 MS Word(COM)로 읽어 번역 docx 생성 (DRM-free 출력)."""
    import comtypes, comtypes.client
    import docx as _docx
    comtypes.CoInitialize()
    word = None
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
        text = d.Content.Text
        d.Close(False)
    finally:
        if word is not None:
            word.Quit()
        comtypes.CoUninitialize()
    # \r(문단), \x07(셀), \x0b(줄바꿈) 등으로 분리
    lines = [ln.strip() for ln in re.split(r'[\r\n\x07\x0b\x0c]', text)]
    paras = [ln for ln in lines if ln]
    job['stage'] = '번역 중'
    cache, lang = translate_unique(paras, job)
    job['stage'] = '문서 작성 중'
    doc = _docx.Document()
    for ln in paras:
        doc.add_paragraph(cache.get(ln, ln))
    doc.save(out_docx)
    return lang


def translate_pdf_to_docx(path, out_docx, job):
    """PDF 본문 텍스트를 추출·번역하여 docx 생성. (DRM PDF는 열 수 없음)"""
    import fitz
    import docx as _docx
    try:
        pdf = fitz.open(path)
    except Exception:
        raise ValueError('PDF를 열 수 없습니다. 회사 DRM으로 보호된 PDF는 번역할 수 없습니다 '
                         '(인가 앱에서 DRM 해제 후 다시 시도).')
    blocks = []
    for page in pdf:
        for b in page.get_text('blocks'):
            t = ' '.join((b[4] or '').split())
            if t:
                blocks.append(t)
    pdf.close()
    if not blocks:
        raise ValueError('PDF에서 텍스트를 찾지 못했습니다(스캔 이미지 PDF일 수 있음).')
    job['stage'] = '번역 중'
    cache, lang = translate_unique([b for b in blocks if LETTER_RE.search(b)], job)
    job['stage'] = '문서 작성 중'
    doc = _docx.Document()
    for b in blocks:
        doc.add_paragraph(cache.get(b, b))
    doc.save(out_docx)
    return lang


def translate_xlsx(path, out_xlsx, job):
    """엑셀(.xlsx/.xlsm)의 문자열 셀만 번역, 수식·숫자·서식은 그대로 보존."""
    import openpyxl
    wb = openpyxl.load_workbook(path)   # 수식 유지(data_only=False)
    cells = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                v = c.value
                if isinstance(v, str) and v.strip() and not v.lstrip().startswith('=') \
                        and LETTER_RE.search(v):
                    cells.append(c)
    job['stage'] = '번역 중'
    cache, lang = translate_unique([c.value for c in cells], job)
    job['stage'] = '엑셀 작성 중'
    for c in cells:
        c.value = cache.get(c.value, c.value)
    wb.save(out_xlsx)
    return lang


def translate_xls_to_xlsx(path, out_xlsx, job):
    """구형 .xls 를 읽어 문자열 셀 번역 후 .xlsx 로 저장(값 기준, 서식 일부 손실)."""
    import xlrd
    from openpyxl import Workbook
    book = xlrd.open_workbook(path)
    wb = Workbook()
    wb.remove(wb.active)
    sheets = []
    strings = set()
    for sh in book.sheets():
        ws = wb.create_sheet(title=(sh.name[:31] or 'Sheet'))
        sheets.append((ws, sh))
        for r in range(sh.nrows):
            for col in range(sh.ncols):
                v = sh.cell_value(r, col)
                if isinstance(v, str) and v.strip() and LETTER_RE.search(v):
                    strings.add(v)
    job['stage'] = '번역 중'
    cache, lang = translate_unique(list(strings), job)
    job['stage'] = '엑셀 작성 중'
    for ws, sh in sheets:
        for r in range(sh.nrows):
            for col in range(sh.ncols):
                v = sh.cell_value(r, col)
                if v == '':
                    continue
                ws.cell(row=r + 1, column=col + 1).value = cache.get(v, v) \
                    if isinstance(v, str) else v
    wb.save(out_xlsx)
    return lang


# ----------------------------- 형식 변환 -----------------------------
def docx_to_pdf(docx_path, pdf_path):
    """MS Word(COM)로 docx -> pdf 변환. Word 설치 필요."""
    import comtypes, comtypes.client
    comtypes.CoInitialize()
    word = None
    try:
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)   # 17 = wdFormatPDF
        doc.Close(False)
    finally:
        if word is not None:
            word.Quit()
        comtypes.CoUninitialize()


def pdf_to_jpegs(pdf_path, out_dir, stem, dpi=150):
    """PDF의 각 페이지를 JPEG 이미지로 저장. 생성된 파일 경로 목록 반환."""
    import fitz
    os.makedirs(out_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    paths = []
    for i in range(doc.page_count):
        pix = doc.load_page(i).get_pixmap(matrix=mat, alpha=False)
        p = os.path.join(out_dir, '%s_KO-%03d.jpg' % (stem, i + 1))
        pix.save(p)
        paths.append(p)
    doc.close()
    return paths


def _to_number(s):
    """문자열이 숫자면 int/float로, 아니면 None."""
    if s is None:
        return None
    t = s.replace(',', '').replace(' ', '').strip()
    if re.fullmatch(r'-?\d+(\.\d+)?', t):
        try:
            return int(t)
        except ValueError:
            try:
                return float(t)
            except ValueError:
                return None
    return None


def _is_total_label(label):
    """합계행 라벨인지 판별. 번역 변형(Total->총/합계/총계 등)까지 포함, '설계' 등 오탐 방지."""
    l = (label or '').strip().lower().replace(' ', '')
    exact = {'계', '총', '합', '합계', '총계', '소계', '총합', '누계', '합산', '총액',
             'total', 'totals', 'sum', 'subtotal', 'grandtotal'}
    if l in exact:
        return True
    return any(l.startswith(w) or l.endswith(w)
               for w in ('합계', '총계', '소계', '총합', '총액', 'total', 'sum', 'subtotal'))


def _is_real_pdf(path):
    """헤더가 %PDF-면 진짜 PDF(=DRM 비암호화). DRM 암호화 시 다른 헤더가 됨."""
    try:
        with open(path, 'rb') as f:
            return f.read(5) == b'%PDF-'
    except Exception:
        return False


def _korean_font():
    for f in (r'C:\Windows\Fonts\malgun.ttf', r'C:\Windows\Fonts\gulim.ttc',
              r'C:\Windows\Fonts\batang.ttc'):
        if os.path.exists(f):
            return f
    return None


def docx_text_lines(docx_path):
    """docx에서 본문 문단/표 텍스트를 줄 목록으로 추출 (이미지 폴백 렌더용)."""
    import docx as _docx
    d = _docx.Document(docx_path)
    lines = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            lines.append(' | '.join(c.text.strip() for c in row.cells))
        lines.append('')
    return lines


def render_text_to_jpegs(lines, out_dir, stem, dpi=150):
    """번역 텍스트를 A4 페이지 이미지(JPEG)로 직접 렌더링 (DRM-free 폴백)."""
    import fitz
    os.makedirs(out_dir, exist_ok=True)
    fp = _korean_font()
    Wp, Hp, margin, fs, leading = 595.0, 842.0, 50.0, 11.0, 16.0
    font = fitz.Font(fontfile=fp) if fp else fitz.Font()
    maxw = Wp - 2 * margin
    wrapped = []
    for ln in lines:
        ln = (ln or '').rstrip('\n')
        if not ln:
            wrapped.append('')
            continue
        cur = ''
        for ch in ln:
            if cur and font.text_length(cur + ch, fs) > maxw:
                wrapped.append(cur)
                cur = ch
            else:
                cur += ch
        wrapped.append(cur)
    per = max(1, int((Hp - 2 * margin) // leading))
    doc = fitz.open()
    for i in range(0, max(1, len(wrapped)), per):
        page = doc.new_page(width=Wp, height=Hp)
        y = margin + fs
        for line in wrapped[i:i + per]:
            if line:
                if fp:
                    page.insert_text((margin, y), line, fontsize=fs,
                                     fontname='kr', fontfile=fp)
                else:
                    page.insert_text((margin, y), line, fontsize=fs)
            y += leading
    mat = fitz.Matrix(dpi / 72.0, dpi / 72.0)
    paths = []
    for idx in range(doc.page_count):
        pix = doc.load_page(idx).get_pixmap(matrix=mat, alpha=False)
        p = os.path.join(out_dir, '%s_KO-%03d.jpg' % (stem, idx + 1))
        pix.save(p)
        paths.append(p)
    doc.close()
    return paths


def _disp_width(s):
    """한글/한자 등 전각 문자는 2칸으로 계산한 표시 폭."""
    w = 0
    for ch in str(s):
        w += 2 if ord(ch) >= 0x1100 else 1
    return w


def docx_tables_to_excel(docx_path, xlsx_path):
    """docx의 표를 가독성 좋은 시트로 추출.
    - 모든 칸에 테두리, 헤더 강조, 자동 열너비/행높이, 줄바꿈, 헤더 고정
    - 숫자는 숫자형으로 저장하고 합계행에는 =SUM 함수를 자동 삽입"""
    import docx as _docx
    from openpyxl import Workbook
    from openpyxl.utils import get_column_letter
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment

    edge = Side(style='thin', color='9CA3AF')
    BORDER = Border(left=edge, right=edge, top=edge, bottom=edge)
    HEAD_FILL = PatternFill('solid', fgColor='4472C4')
    HEAD_FONT = Font(bold=True, color='FFFFFF', size=11)
    ALT_FILL = PatternFill('solid', fgColor='EAF0FA')
    TOTAL_FILL = PatternFill('solid', fgColor='FCE4D6')
    TOTAL_FONT = Font(bold=True)
    WRAP = Alignment(wrap_text=True, vertical='center', horizontal='left')
    WRAP_C = Alignment(wrap_text=True, vertical='center', horizontal='center')
    NUM_AL = Alignment(wrap_text=True, vertical='center', horizontal='right')

    def style_sheet(ws, grid, ncols, has_header=True):
        # 열 너비 (내용 기반, 10~55칸)
        for j in range(ncols):
            mx = max((_disp_width(grid[i][j]) for i in range(len(grid))
                      if j < len(grid[i])), default=8)
            ws.column_dimensions[get_column_letter(j + 1)].width = min(55, max(10, mx + 3))
        # 셀 서식 + 행 높이
        for i in range(len(grid)):
            is_total = bool(grid[i]) and _is_total_label(grid[i][0])
            max_lines = 1
            for j in range(ncols):
                cell = ws.cell(row=i + 1, column=j + 1)
                cell.border = BORDER
                if has_header and i == 0:
                    cell.fill = HEAD_FILL
                    cell.font = HEAD_FONT
                    cell.alignment = WRAP_C
                elif is_total:
                    cell.fill = TOTAL_FILL
                    cell.font = TOTAL_FONT
                    cell.alignment = NUM_AL if isinstance(cell.value, (int, float)) \
                        or (isinstance(cell.value, str) and cell.value.startswith('=')) else WRAP
                else:
                    if i % 2 == 0:
                        cell.fill = ALT_FILL
                    cell.alignment = NUM_AL if isinstance(cell.value, (int, float)) else WRAP
                # 줄 수 추정(행 높이용)
                colw = ws.column_dimensions[get_column_letter(j + 1)].width or 10
                txt = '' if cell.value is None else str(cell.value)
                lines = max(txt.count('\n') + 1,
                            -(-_disp_width(txt) // max(1, int(colw - 1))))
                max_lines = max(max_lines, lines)
            ws.row_dimensions[i + 1].height = min(120, max(18, max_lines * 15.5))
        if has_header:
            ws.freeze_panes = 'A2'

    d = _docx.Document(docx_path)
    wb = Workbook()
    wb.remove(wb.active)
    tnum = 0
    for table in d.tables:
        grid = [[c.text.strip() for c in row.cells] for row in table.rows]
        if not grid or not any(any(c for c in r) for r in grid):
            continue
        tnum += 1
        ws = wb.create_sheet(title='표%d' % tnum)
        ncols = max(len(r) for r in grid)
        for i, row in enumerate(grid):
            for j in range(ncols):
                val = row[j] if j < len(row) else ''
                num = _to_number(val)
                ws.cell(row=i + 1, column=j + 1).value = num if num is not None else val
        # 합계행: 숫자 열에 =SUM 삽입
        for i, row in enumerate(grid):
            if not row or not _is_total_label(row[0]):
                continue
            for j in range(1, ncols):
                if any(j < len(grid[k]) and _to_number(grid[k][j]) is not None
                       for k in range(i)):
                    col = get_column_letter(j + 1)
                    ws.cell(row=i + 1, column=j + 1).value = '=SUM(%s1:%s%d)' % (col, col, i)
        style_sheet(ws, grid, ncols, has_header=True)

    if tnum == 0:
        # 표가 없는 문서: 제목/본문을 구분되게 한 열로 정리
        from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
        edge2 = Side(style='thin', color='D0D0D0')
        b2 = Border(left=edge2, right=edge2, top=edge2, bottom=edge2)
        ws = wb.create_sheet(title='내용')
        ws.column_dimensions['A'].width = 95
        ws.cell(row=1, column=1).value = '문서 내용 (표 없음 — 문단별 정리)'
        ws.cell(row=1, column=1).font = Font(bold=True, color='FFFFFF')
        ws.cell(row=1, column=1).fill = PatternFill('solid', fgColor='4472C4')
        r = 2
        for para in d.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            cell = ws.cell(row=r, column=1)
            cell.value = t
            cell.border = b2
            cell.alignment = Alignment(wrap_text=True, vertical='center')
            is_head = 'head' in (para.style.name or '').lower() or 'title' in (para.style.name or '').lower()
            if is_head:
                cell.font = Font(bold=True, size=12, color='1F3864')
                cell.fill = PatternFill('solid', fgColor='DCE6F1')
            elif r % 2 == 0:
                cell.fill = PatternFill('solid', fgColor='F4F7FB')
            lines = -(-_disp_width(t) // 92)
            ws.row_dimensions[r].height = min(150, max(18, lines * 15.5))
            r += 1
        ws.freeze_panes = 'A2'

    wb.save(xlsx_path)


# ----------------------------- 웹 UI -----------------------------
PAGE = """<!doctype html><html lang="ko"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>자동 번역기 → 한국어</title>
<style>
 *{box-sizing:border-box;font-family:'Segoe UI','맑은 고딕',sans-serif}
 body{margin:0;background:#0f172a;color:#e2e8f0;display:flex;justify-content:center;padding:40px 16px}
 .card{width:100%;max-width:680px;background:#1e293b;border-radius:16px;padding:32px;box-shadow:0 10px 40px rgba(0,0,0,.4)}
 h1{margin:0 0 6px;font-size:22px}
 p.sub{margin:0 0 24px;color:#94a3b8;font-size:14px}
 label{display:block;font-size:13px;color:#cbd5e1;margin-bottom:8px}
 input[type=text]{width:100%;padding:13px 14px;border-radius:10px;border:1px solid #334155;background:#0f172a;color:#e2e8f0;font-size:14px}
 button{margin-top:16px;width:100%;padding:14px;border:0;border-radius:10px;background:#3b82f6;color:#fff;font-size:15px;font-weight:600;cursor:pointer}
 button:disabled{background:#475569;cursor:not-allowed}
 .bar{height:10px;background:#0f172a;border-radius:6px;overflow:hidden;margin-top:20px;display:none}
 .bar>i{display:block;height:100%;width:0;background:#22c55e;transition:width .3s}
 #log{margin-top:16px;font-size:13px;line-height:1.7;white-space:pre-wrap;word-break:break-all}
 .ok{color:#4ade80}.err{color:#f87171}.muted{color:#94a3b8}
 .hint{margin-top:18px;font-size:12px;color:#64748b;border-top:1px solid #334155;padding-top:14px}
</style></head><body>
<div class="card">
 <h1>🌐 자동 번역기 → 한국어</h1>
 <p class="sub">파일 경로를 입력하면 언어를 자동 감지해 한국어로 번역하고, <b>같은 폴더</b>에 <b>_KO</b> 파일로 저장합니다.</p>
 <label>파일 전체 경로 (탐색기에서 파일 우클릭 → "경로로 복사")</label>
 <input id="path" type="text" placeholder="예: D:\\문서\\manual.docx" autofocus>
 <button id="go" onclick="start()">번역 시작</button>
 <div class="bar" id="bar"><i id="fill"></i></div>
 <div id="log"></div>
 <div class="hint">지원 형식: .docx (서식·이미지·표 유지), .txt &nbsp;|&nbsp; 번역 엔진: Google (인터넷 필요)</div>
</div>
<script>
let timer=null;
function log(msg,cls){document.getElementById('log').innerHTML='<span class="'+(cls||'')+'">'+msg+'</span>';}
async function start(){
 const path=document.getElementById('path').value.trim();
 if(!path){log('파일 경로를 입력하세요.','err');return;}
 const btn=document.getElementById('go');btn.disabled=true;
 document.getElementById('bar').style.display='block';
 document.getElementById('fill').style.width='0%';
 log('작업을 시작합니다...','muted');
 let r=await fetch('/start',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({path})});
 let j=await r.json();
 if(j.error){log('오류: '+j.error,'err');btn.disabled=false;return;}
 const id=j.id;
 timer=setInterval(async()=>{
   let s=await(await fetch('/status?id='+id)).json();
   let pct=s.total?Math.round(s.done/s.total*100):0;
   document.getElementById('fill').style.width=pct+'%';
   if(s.status==='running'){
     log(s.stage+' ... '+(s.total?(s.done+' / '+s.total+' 문단 ('+pct+'%)'):''),'muted');
   }else if(s.status==='done'){
     clearInterval(timer);btn.disabled=false;
     document.getElementById('fill').style.width='100%';
     log('✅ 번역 완료!\\n감지된 언어: '+s.lang+'\\n저장 위치:\\n'+s.output,'ok');
   }else if(s.status==='error'){
     clearInterval(timer);btn.disabled=false;
     log('❌ 오류: '+s.error,'err');
   }
 },600);
}
document.getElementById('path').addEventListener('keydown',e=>{if(e.key==='Enter')start();});
</script></body></html>"""


class Handler(BaseHTTPRequestHandler):
    def _send(self, code, body, ctype='application/json'):
        data = body.encode('utf-8') if isinstance(body, str) else body
        self.send_response(code)
        self.send_header('Content-Type', ctype + '; charset=utf-8')
        self.send_header('Content-Length', str(len(data)))
        # index.html 을 파일(file://)로 직접 열어도 서버와 통신되도록 허용
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
        self.wfile.write(data)

    def do_OPTIONS(self):
        # 브라우저 CORS 프리플라이트 응답
        self.send_response(204)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, GET, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.send_header('Content-Length', '0')
        self.end_headers()

    def do_GET(self):
        if self.path == '/' or self.path.startswith('/index'):
            # 같은 폴더에 index.html 이 있으면 그 파일을, 없으면 내장 페이지를 보여줌
            html_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                     'index.html')
            try:
                page = open(html_path, 'r', encoding='utf-8').read()
            except Exception:
                page = PAGE
            self._send(200, page, 'text/html')
        elif self.path.startswith('/status'):
            q = urllib.parse.urlparse(self.path).query
            jid = urllib.parse.parse_qs(q).get('id', [''])[0]
            job = JOBS.get(jid)
            if not job:
                self._send(404, json.dumps({'error': 'no job'}))
            else:
                self._send(200, json.dumps({
                    'status': job['status'], 'stage': job.get('stage', ''),
                    'done': job.get('done', 0), 'total': job.get('total', 0),
                    'output': job.get('output', ''), 'lang': job.get('lang', ''),
                    'error': job.get('error', '')}))
        else:
            self._send(404, json.dumps({'error': 'not found'}))

    def do_POST(self):
        if self.path == '/start':
            ln = int(self.headers.get('Content-Length', 0))
            body = json.loads(self.rfile.read(ln).decode('utf-8'))
            path = body.get('path', '')
            formats = body.get('formats') or ['docx']
            jid = uuid.uuid4().hex
            with JOBS_LOCK:
                JOBS[jid] = {'status': 'running', 'stage': '대기 중',
                             'done': 0, 'total': 0}
            threading.Thread(target=run_job, args=(jid, path, formats), daemon=True).start()
            self._send(200, json.dumps({'id': jid}))
        else:
            self._send(404, json.dumps({'error': 'not found'}))

    def log_message(self, *a):
        pass   # 콘솔 로그 끔


def cli_translate(input_str, status_file=None, formats=None):
    """명령줄/HTA용 단일 번역 실행. 진행상태를 status_file(JSON)에 계속 기록."""
    job = {'status': 'running', 'stage': '준비 중', 'done': 0, 'total': 0}

    def dump():
        if not status_file:
            return
        try:
            with open(status_file, 'w', encoding='utf-8') as f:
                json.dump({'status': job['status'], 'stage': job.get('stage', ''),
                           'done': job.get('done', 0), 'total': job.get('total', 0),
                           'output': job.get('output', ''), 'lang': job.get('lang', ''),
                           'error': job.get('error', '')}, f, ensure_ascii=True)
        except Exception:
            pass

    stop = threading.Event()

    def writer():
        while not stop.is_set():
            dump()
            time.sleep(0.4)
        dump()

    th = threading.Thread(target=writer, daemon=True)
    th.start()
    JOBS['cli'] = job
    run_job('cli', input_str, formats)
    stop.set()
    th.join(timeout=2)
    dump()
    if job['status'] == 'done':
        print('RESULT:' + job.get('output', ''))
        print('LANG:' + str(job.get('lang', '')))
        return 0
    else:
        print('ERROR:' + job.get('error', ''))
        return 1


def _write_status(status_file, **kw):
    if not status_file:
        return
    base = {'status': 'running', 'stage': '', 'done': 0, 'total': 0,
            'output': '', 'lang': '', 'error': ''}
    base.update(kw)
    try:
        with open(status_file, 'w', encoding='utf-8') as f:
            json.dump(base, f, ensure_ascii=True)
    except Exception:
        pass


def ensure_deps(status_file=None):
    """필요한 외부 패키지(lxml, python-docx, bs4)를 현재 파이썬에 자동 설치(없을 때만).
    어떤 파이썬으로 실행되든 스스로 의존성을 갖추도록 한다."""
    import importlib, subprocess
    for mod, pkg in [('lxml', 'lxml'),
                     ('docx', 'python-docx'),
                     ('bs4', 'beautifulsoup4'),
                     ('comtypes', 'comtypes'),     # PDF 변환/.doc 읽기(Word COM)
                     ('fitz', 'PyMuPDF'),          # JPEG 변환/.pdf 텍스트 추출
                     ('openpyxl', 'openpyxl'),     # Excel 읽기/생성
                     ('xlrd', 'xlrd')]:            # 구형 .xls 읽기
        try:
            importlib.import_module(mod)
            continue
        except ImportError:
            pass
        _write_status(status_file, stage='필요 패키지 설치 중 (%s) ...' % pkg)
        try:
            subprocess.check_call([sys.executable, '-m', 'pip', 'install',
                                   '--quiet', '--disable-pip-version-check', pkg])
            importlib.invalidate_caches()
        except Exception:
            pass  # 설치 실패 시 이후 단계에서 명확한 오류로 안내됨


def main():
    # 콘솔 인코딩(cp949)에서 못 쓰는 문자가 있어도 죽지 않도록
    try:
        sys.stdout.reconfigure(errors='backslashreplace')
        sys.stderr.reconfigure(errors='backslashreplace')
    except Exception:
        pass

    # CLI 모드:  python auto_translator.py --cli "<경로나URL>" "<상태파일>" "docx,pdf,jpeg,excel"
    if len(sys.argv) >= 3 and sys.argv[1] == '--cli':
        inp = sys.argv[2]
        stf = sys.argv[3] if len(sys.argv) >= 4 else None
        fmts = sys.argv[4].split(',') if len(sys.argv) >= 5 else ['docx']
        _write_status(stf, stage='준비 중...')   # HTA가 곧바로 상태를 읽도록
        ensure_deps(stf)                          # 필요 패키지 자동 설치
        raise SystemExit(cli_translate(inp, stf, fmts))

    ensure_deps()                                 # 서버 모드도 의존성 자동 확보
    srv = ThreadingHTTPServer(('127.0.0.1', PORT), Handler)
    url = 'http://127.0.0.1:%d/' % PORT
    print('=' * 50)
    print(' 자동 번역기 실행 중')
    print(' 브라우저 주소:', url)
    print(' 종료하려면 이 창에서 Ctrl + C')
    print('=' * 50)
    try:
        webbrowser.open(url)
    except Exception:
        pass
    try:
        srv.serve_forever()
    except KeyboardInterrupt:
        print('\n종료합니다.')


if __name__ == '__main__':
    main()
